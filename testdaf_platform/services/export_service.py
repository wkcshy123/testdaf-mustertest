"""题库导出服务 — Word (.docx)，保存到用户 Downloads 目录。"""

import re
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from testdaf_platform.storage.question_bank import QuestionBank

DOWNLOADS_DIR = Path.home() / "Downloads"

_SECTION_LABELS = {
    "listening": "Hoerverstehen",
    "reading": "Leseverstehen",
    "writing": "Schriftlicher Ausdruck",
    "speaking": "Muendlicher Ausdruck",
}


def safe_filename(title: str, max_len: int = 50) -> str:
    name = re.sub(r'[<>:\"/\\|?*]', "", title)
    name = re.sub(r"\s+", "_", name.strip())
    if len(name) <= max_len:
        return name
    words = re.findall(r"[A-Za-z\u00C4\u00D6\u00DC\u00E4\u00F6\u00FC\u00DF]{3,}", title)
    result = "_".join(words)
    if len(result) > max_len:
        result = result[:max_len].rstrip("_")
    return result if result else name[:max_len]


def _stripped(value: object) -> str:
    return str(value or "").strip()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"


def _add_para(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def _add_bold_para(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True


def _add_highlighted(doc: Document, label: str, value: str) -> None:
    para = doc.add_paragraph()
    rl = para.add_run(f"{label} ")
    rl.font.name = "Arial"
    rl.font.size = Pt(11)
    rl.bold = True
    rv = para.add_run(value)
    rv.font.name = "Arial"
    rv.font.size = Pt(11)
    rpr = rv._element.get_or_add_rPr()
    hl = rpr.makeelement(qn("w:highlight"), {qn("w:val"): "yellow"})
    rpr.append(hl)


def _save_to_downloads(doc: Document, title: str) -> Path:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
    name = f"{safe_filename(title)}_{uuid4().hex[:6]}.docx"
    filepath = DOWNLOADS_DIR / name
    doc.save(str(filepath))
    return filepath


def _add_docx_header(doc: Document, section: str, task_type: str) -> None:
    label = task_type.replace("_", " ").title()
    h = doc.add_heading(f"{section}  —  {label}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para(doc, "")


class ExportService:

    def __init__(self, bank: QuestionBank | None = None):
        self.bank = bank or QuestionBank()

    # ==================================================================
    # 自动保存 — 生成时调用，含答案，答案黄底高亮
    # ==================================================================

    def save_reading_with_answers(self, generation: dict, task_type: str) -> Path:
        methods = {
            "aufgabe_1": self._docx_reading_a1_answers,
            "aufgabe_2": self._docx_reading_a2_answers,
            "aufgabe_3": self._docx_reading_a3_answers,
        }
        method = methods.get(task_type)
        if method is None:
            raise ValueError(f"不支持的阅读题型: {task_type}")
        doc = Document()
        _add_docx_header(doc, _SECTION_LABELS.get("reading", "Reading"), task_type)
        method(doc, generation)
        return _save_to_downloads(doc, generation.get("title", task_type))

    def save_listening_with_answers(self, generation: dict, task_type: str) -> Path:
        methods = {
            "aufgabe_1": self._docx_listening_a1_answers,
            "aufgabe_2": self._docx_listening_a2_answers,
            "aufgabe_3": self._docx_listening_a3_answers,
        }
        method = methods.get(task_type)
        if method is None:
            raise ValueError(f"不支持的听力题型: {task_type}")
        doc = Document()
        _add_docx_header(doc, _SECTION_LABELS.get("listening", "Listening"), task_type)
        method(doc, generation)
        return _save_to_downloads(doc, generation.get("title", task_type))

    # ==================================================================
    # 手动下载 — 仅题目，无答案
    # ==================================================================

    def export_questions_only(self, relative_path: str) -> Path:
        bundle = self.bank.load_question_bundle(relative_path)
        manifest = bundle.get("manifest", {})
        section = manifest.get("section", "")
        task_type = manifest.get("task_type", "")
        title = manifest.get("title", "Export")

        doc = Document()
        _add_docx_header(doc, _SECTION_LABELS.get(section, section), task_type)

        if section == "reading":
            self._docx_reading_questions(doc, bundle, task_type)
        elif section == "listening":
            self._docx_listening_questions(doc, bundle)
        elif section == "writing":
            self._docx_writing_questions(doc, bundle)
        elif section == "speaking":
            self._docx_speaking_questions(doc, bundle)
        else:
            raise ValueError(f"不支持的科目: {section}")
        return _save_to_downloads(doc, title)

    # ==================================================================
    # Reading 含答案
    # ==================================================================

    def _docx_reading_a1_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A1")), level=1)
        _add_para(doc, _stripped(g.get("topic", "")))

        _add_heading(doc, "Aufgaben", level=2)
        for a in g.get("answers", []):
            _add_para(doc, f"{a.get('number', '?')}.  Welches Angebot passt?")

        _add_heading(doc, "Anzeigen", level=2)
        for o in g.get("offers", []):
            _add_para(doc, f"[{_stripped(o.get('label', ''))}] {_stripped(o.get('text', ''))}")

        _add_heading(doc, "Profile", level=2)
        for p in g.get("profiles", []):
            _add_para(doc, f"[{_stripped(p.get('label', ''))}] {_stripped(p.get('description', ''))}")

        _add_heading(doc, "Antworten", level=2)
        for a in g.get("answers", []):
            n = a.get("number", "?")
            corr = f"Angebot {_stripped(a.get('correct_offer_label', ''))} → Profil {_stripped(a.get('correct_profile_label', ''))}"
            _add_highlighted(doc, f"{n}. Antwort:", corr)
            expl = _stripped(a.get("explanation", ""))
            if expl:
                _add_highlighted(doc, "   Erklaerung:", expl)

    def _docx_reading_a2_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A2")), level=1)
        _add_para(doc, _stripped(g.get("topic", "")))

        _add_heading(doc, "Aufgaben", level=2)
        for q in g.get("questions", []):
            _add_para(doc, f"{q.get('number', '?')}.  {_stripped(q.get('prompt', ''))}")
            for o in q.get("options", []):
                _add_para(doc, f"     {_stripped(o.get('label', ''))}) {_stripped(o.get('text', ''))}")

        _add_heading(doc, "Lesetext", level=2)
        _add_para(doc, _stripped(g.get("reading_text", "")))

        _add_heading(doc, "Antworten", level=2)
        for q in g.get("questions", []):
            n = q.get("number", "?")
            ans = q.get("answer", {})
            ans = ans.get("label", "") if isinstance(ans, dict) else ans
            _add_highlighted(doc, f"{n}. Antwort:", _stripped(ans))
            expl = _stripped(q.get("explanation", ""))
            if expl:
                _add_highlighted(doc, "   Erklaerung:", expl)

    def _docx_reading_a3_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A3")), level=1)
        _add_para(doc, _stripped(g.get("topic", "")))

        _add_heading(doc, "Aufgaben", level=2)
        for s in g.get("statements", []):
            _add_para(doc, f"{s.get('number', '?')}.  {_stripped(s.get('statement', ''))}")
            _add_para(doc, "     ☐ Ja   ☐ Nein   ☐ Text sagt dazu nichts")

        _add_heading(doc, "Lesetext", level=2)
        _add_para(doc, _stripped(g.get("reading_text", "")))

        _add_heading(doc, "Antworten", level=2)
        for s in g.get("statements", []):
            n = s.get("number", "?")
            _add_highlighted(doc, f"{n}. Antwort:", _stripped(s.get("answer", "")))
            expl = _stripped(s.get("explanation", ""))
            if expl:
                _add_highlighted(doc, "   Erklaerung:", expl)

    # ==================================================================
    # Listening 含答案
    # ==================================================================

    def _docx_listening_a1_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A1")), level=1)
        _add_para(doc, f"Thema: {_stripped(g.get('topic', ''))}")
        _add_para(doc, f"Beziehung: {_stripped(g.get('relationship', ''))}")

        _add_heading(doc, "Aufgaben", level=2)
        for q in g.get("questions", []):
            _add_para(doc, f"{q.get('number', '?')}.  {_stripped(q.get('prompt', ''))}")

        _add_heading(doc, "Transkript", level=2)
        _add_para(doc, _stripped(g.get("transcript", "")))

        _add_heading(doc, "Antworten", level=2)
        for q in g.get("questions", []):
            n = q.get("number", "?")
            ans = q.get("answer", [])
            if isinstance(ans, list):
                ans = "; ".join(_stripped(a) for a in ans)
            else:
                ans = _stripped(ans)
            _add_highlighted(doc, f"{n}. Antwort:", ans)
            ev = _stripped(q.get("evidence", ""))
            if ev:
                _add_highlighted(doc, "   Beleg:", ev)

    def _docx_listening_a2_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A2")), level=1)
        _add_para(doc, f"Thema: {_stripped(g.get('topic', ''))}")
        _add_para(doc, f"Format: {_stripped(g.get('format_note', ''))}")

        sr = g.get("speaker_roles", {})
        if sr:
            _add_para(doc, "Sprecher:")
            for sid, role in sr.items():
                _add_para(doc, f"  • {_stripped(sid)} — {_stripped(role)}")

        _add_heading(doc, "Aufgaben", level=2)
        for s in g.get("statements", []):
            _add_para(doc, f"{s.get('number', '?')}.  {_stripped(s.get('statement', ''))}")
            _add_para(doc, "     ☐ Richtig   ☐ Falsch")

        _add_heading(doc, "Transkript", level=2)
        _add_para(doc, _stripped(g.get("transcript", "")))

        _add_heading(doc, "Antworten", level=2)
        for s in g.get("statements", []):
            n = s.get("number", "?")
            _add_highlighted(doc, f"{n}. Antwort:", _stripped(s.get("answer", "")))
            ev = _stripped(s.get("evidence", ""))
            if ev:
                _add_highlighted(doc, "   Beleg:", ev)
            dt = _stripped(s.get("distractor_explanation", ""))
            if dt:
                _add_highlighted(doc, "   Distraktor:", dt)

    def _docx_listening_a3_answers(self, doc: Document, g: dict) -> None:
        _add_heading(doc, _stripped(g.get("title", "A3")), level=1)
        _add_para(doc, f"Thema: {_stripped(g.get('topic', ''))}")
        _add_para(doc, f"Format: {_stripped(g.get('format_note', ''))}")

        sr = g.get("speaker_roles", {})
        if sr:
            _add_para(doc, "Sprecher:")
            for sid, role in sr.items():
                _add_para(doc, f"  • {_stripped(sid)} — {_stripped(role)}")

        _add_heading(doc, "Aufgaben", level=2)
        for q in g.get("questions", []):
            _add_para(doc, f"{q.get('number', '?')}.  {_stripped(q.get('prompt', ''))}")

        _add_heading(doc, "Transkript", level=2)
        _add_para(doc, _stripped(g.get("transcript", "")))

        _add_heading(doc, "Antworten", level=2)
        for q in g.get("questions", []):
            n = q.get("number", "?")
            ans = q.get("answer", [])
            if isinstance(ans, list):
                ans = "; ".join(_stripped(a) for a in ans)
            else:
                ans = _stripped(ans)
            _add_highlighted(doc, f"{n}. Antwort:", ans)
            ev = _stripped(q.get("evidence", ""))
            if ev:
                _add_highlighted(doc, "   Beleg:", ev)

    # ==================================================================
    # 仅题目版
    # ==================================================================

    def _docx_reading_questions(self, doc: Document, bundle: dict, task_type: str) -> None:
        if task_type == "aufgabe_1":
            self._docx_reading_a1_questions(doc, bundle)
        elif task_type == "aufgabe_2":
            self._docx_reading_a2_questions(doc, bundle)
        elif task_type == "aufgabe_3":
            self._docx_reading_a3_questions(doc, bundle)

    def _docx_reading_a1_questions(self, doc: Document, bundle: dict) -> None:
        profiles = bundle.get("profiles", [])
        offers = bundle.get("texts", [])
        example_label = bundle.get("manifest", {}).get("parameters", {}).get("example_offer_label")

        doc.add_heading("Aufgabenstellung", level=2)
        _add_para(doc, "Sie suchen fuer verschiedene Personen eine passende Anzeige. "
                    "Schreiben Sie den jeweiligen Buchstaben (A–H). "
                    "Wenn es keine passende Anzeige gibt, schreiben Sie I.")

        doc.add_heading("Personen 1 – 10", level=2)
        for p in profiles:
            _add_para(doc, f"{p['number']}.  {p.get('need', '')}")

        doc.add_heading("Texte A – H", level=2)
        for o in offers:
            if example_label and o.get("label") == example_label:
                continue
            _add_bold_para(doc, f"{o['label']}.  {o.get('heading', '')}")
            _add_para(doc, o.get("text", ""))

    def _docx_reading_a2_questions(self, doc: Document, bundle: dict) -> None:
        text = bundle.get("reading_text", "")
        paragraphs = bundle.get("paragraphs", [])
        questions = bundle.get("questions", [])

        if paragraphs:
            for p_item in paragraphs:
                _add_para(doc, p_item.get("text", ""))
                _add_para(doc, "")
        elif text:
            _add_para(doc, text)

        doc.add_heading("Aufgaben 11 – 20", level=2)
        for q in questions:
            _add_bold_para(doc, f"{q['number']}.  {q.get('prompt', '')}")
            opts = q.get("options", {})
            for label in ("A", "B", "C"):
                _add_para(doc, f"    {label}    {opts.get(label, '')}")
            _add_para(doc, "")

    def _docx_reading_a3_questions(self, doc: Document, bundle: dict) -> None:
        text = bundle.get("reading_text", "")
        paragraphs = bundle.get("paragraphs", [])
        stmts = bundle.get("statements", []) or bundle.get("questions", [])

        if paragraphs:
            for p_item in paragraphs:
                _add_para(doc, p_item.get("text", ""))
                _add_para(doc, "")
        elif text:
            _add_para(doc, text)

        doc.add_heading("Aufgaben 21 – 30: Ja / Nein / Text sagt dazu nichts", level=2)
        for s in stmts:
            _add_para(doc, f"{s.get('number', '')}.  {s.get('statement', '')}")
            _add_para(doc, "     ☐ Ja   ☐ Nein   ☐ Text sagt dazu nichts")

    def _docx_listening_questions(self, doc: Document, bundle: dict) -> None:
        qs = bundle.get("questions", []) or bundle.get("statements", [])

        doc.add_heading("Aufgaben", level=2)
        for q in qs:
            num = q.get("number", "")
            prompt = q.get("prompt") or q.get("statement", "")
            _add_para(doc, f"{num}.  {prompt}")

        if bundle.get("transcript"):
            doc.add_heading("Transkript", level=2)
            _add_para(doc, bundle["transcript"])

    def _docx_writing_questions(self, doc: Document, bundle: dict) -> None:
        pd = bundle.get("prompt", {})
        bg = pd.get("background", "") or pd.get("topic", "")
        tp = pd.get("task_prompt", "")
        ins = pd.get("writing_instructions", [])

        if bg:
            _add_para(doc, bg)
            _add_para(doc, "")
        doc.add_heading("Aufgabe", level=2)
        _add_bold_para(doc, tp)
        _add_para(doc, "")
        if ins:
            doc.add_heading("Bearbeiten Sie die folgenden Punkte:", level=2)
            for i in ins:
                _add_para(doc, f"- {i}")

        cs = bundle.get("charts", [])
        if cs:
            doc.add_heading("Grafiken", level=2)
            for idx, c in enumerate(cs, 1):
                doc.add_heading(f"Grafik {idx}: {c.get('title', '')}", level=3)
                data = c.get("data", [])
                if data:
                    for item in data:
                        _add_para(doc, f"  {item.get('label', '')}: {item.get('value', '')}")

    def _docx_speaking_questions(self, doc: Document, bundle: dict) -> None:
        manifest = bundle.get("manifest", {})
        if manifest.get("task_type") == "test_set":
            tasks = bundle.get("tasks", [])
            for i, task in enumerate(tasks):
                self._write_speaking_task(doc, task)
                if i < len(tasks) - 1:
                    doc.add_page_break()
        else:
            self._write_speaking_task(doc, bundle.get("prompt", {}))

    def _write_speaking_task(self, doc: Document, task: dict) -> None:
        num = task.get("number", "")
        tl = task.get("task_type", "")
        prep = task.get("prep_time", "")
        speak = task.get("speaking_time", "")
        sc = task.get("scenario", "")
        pts = task.get("prompt_points", [])
        er = task.get("examiner_role", "")

        doc.add_heading(f"Aufgabe {num}: {tl}", level=2)
        line = f"Vorbereitungszeit: {prep}"
        if speak:
            line += f"  |  Sprechzeit: {speak}"
        if er:
            line += f"  |  Gespraechspartner/in: {er}"
        _add_para(doc, line)
        doc.add_heading("Situation", level=3)
        _add_para(doc, sc)
        doc.add_heading("Ihre Aufgabe", level=3)
        for pt in pts:
            _add_para(doc, f"- {pt}")
